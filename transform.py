from PyQt5.QtWidgets import QWidget, QPlainTextEdit, QTextEdit, QPushButton, QGridLayout, QMessageBox, QFileDialog, QLabel
from PyQt5.QtCore import pyqtSignal, QTimer
from PyQt5.QtGui import QTextCursor
from voice_transform import VoiceTransThread
from docx import Document
from docx.oxml.ns import qn
from file import File
from pandas import DataFrame
import os
import re


class VoiceTans(QWidget):
    transform_status = pyqtSignal(list)
    def __init__(self, parent = None):
        super(VoiceTans, self).__init__()
        self.__files = []
        self.cur_show2text_idx = 0
        self.cur_playing_postion = 0
        self.playing_file_idx = None
        self.__init_layout()
        self.voice_trans_thread = None
        self.timer = QTimer()
        self.timer.timeout.connect(self.updata_file_list_status)
        self.setParent(parent)

    def __init_layout(self):
        # 创建文本区
        self.__text = QPlainTextEdit(self)
        self.__text.setReadOnly(True)
        self.__text.setMinimumSize(600, 300)

        # 创建语音对应修改的文本框
        self.__text_edit = QTextEdit(self)
        self.__text_edit.setMinimumSize(500,100)

        #创建保存修改按键
        self.__btn_save = QPushButton('保存修改内容', self)
        self.__btn_save.setMaximumSize(50, 30)
        self.__btn_save.clicked.connect(self.on_btn_save_clicked)

        # 创建布局管理器管理控件
        gridlayout = QGridLayout(self)
        gridlayout.setRowStretch(1, 6)
        gridlayout.setRowStretch(3, 1)
        gridlayout.addWidget(QLabel("语音转写区:"), 0 , 0, 1, 2)
        gridlayout.addWidget(self.__text, 1, 0, 1, 2)
        gridlayout.addWidget(QLabel('文本修改区:'), 2, 0)
        gridlayout.addWidget(self.__text_edit, 3, 0)
        gridlayout.addWidget(self.__btn_save, 3, 1)

        self.setLayout(gridlayout)

    # 是否存在没转写的文件
    def exit_files2trasn(self):
        for file in self.__files:
            if not file.file_status == 'Success':
                return True
        return False

    def change_playing_idx(self, index):
        self.__text_edit.clear()
        self.__text.clear()
        self.playing_file_idx = index

    # 保存按钮槽函数
    def on_btn_save_clicked(self):
        if not self.__files:
            QMessageBox.information(self, '警告', '请先打开文件')
            return
        ret = QMessageBox.information(self, '警告', '是否保存文件', QMessageBox.No|QMessageBox.Yes)
        if ret == QMessageBox.Yes:
            tmp_str = ''
            for slice_txt in self.__files[self.cur_show2text_idx].voice_msg:
                if not self.cur_playing_postion:
                   return
                if (self.cur_playing_postion >= slice_txt['text_begin']) and\
                   (self.cur_playing_postion <= slice_txt['text_end']):
                    slice_txt['text'] = self.__text_edit.toPlainText()
                tmp_str += slice_txt['text']

            self.__files[self.cur_show2text_idx].set_file_txt(tmp_str)
            self.__text.clear()
            self.__text.appendPlainText(self.__files[self.cur_show2text_idx].get_file_txt())

    # 语音转写功能按钮槽函数
    def on_btn_voice_trans_clicked(self, index):
        if not self.__files:
            QMessageBox(QMessageBox.Warning, '警告', '请打开要转写的语音文件').exec()
        else:
            if self.exit_files2trasn() and (not self.voice_trans_thread):
                self.voice_trans_thread = VoiceTransThread(self.__files)
                self.voice_trans_thread.trans_end.connect(self.voice_trans_end)
                self.timer.start(200)
                self.voice_trans_thread.start()

    # 语音转写完成之后退出线程
    def voice_trans_end(self, thread_back):
        if isinstance(thread_back[0], str):
            QMessageBox(QMessageBox.Warning, '警告', ''.join(thread_back)).exec()
        elif isinstance(thread_back[0], File):
            # 设置各个控件的状态
            self.__text.appendPlainText(self.__files[0].get_file_txt())

        # 退出线程
        self.voice_trans_thread.quit()
        self.transform_status.emit(self.__files)
        self.timer.stop()
        self.voice_trans_thread = None

    def updata_file_list_status(self):
        self.transform_status.emit(self.__files)

    def stop2show_playing_file(self):
        self.__text.clear()
        self.__text_edit.clear()

    def on_playing_show(self, position):
        def do_hilght_search(target):
            text = self.__text.toPlainText()
            index = text.find(target[:5])
            if index > 0:
                cursor = self.__text.textCursor()
                cursor.setPosition(index)
                cursor.setPosition(index + len(target), QTextCursor.KeepAnchor)
                self.__text.setTextCursor(cursor)
                self.__text.repaint()

        if not self.__files:
            return
        if self.__files[self.playing_file_idx]:
            slice_text = self.__files[self.playing_file_idx].voice_msg
            if slice_text:
                self.cur_playing_postion = position
                for dict in slice_text:
                    if (position >= dict['text_begin']) and (position <= dict['text_end']):
                        do_hilght_search(dict['text'])
                        if self.__text_edit.toPlainText() != dict['text']:
                            self.__text_edit.clear()
                            self.__text_edit.setText(dict['text'])
                            self.__text_edit.repaint()

    def get_open_files(self, file):
        # 如果发现语音转写线程被开启，则把后面加上文件状态改为正在转写状态
        if self.voice_trans_thread:
            file.file_status = 'Running'
        self.__files.append(file)

    # 点击文件列表的时候，显示选中文件的内容
    def show_file_txt(self, index):
        self.cur_show2text_idx = index
        # 获取被选中的文件
        self.__files[index]
        self.__text.clear()
        # 如果文件完成转写，则显示转写的内容，如果还在转写，则显示正在转写中
        if self.__files[index].file_status == 'Success':
            self.__text.appendPlainText(self.__files[index].get_file_txt())
            self.__text.repaint()

    # 转存语音文件按键点击槽函数
    def on_btn_export_clicked(self):
        flag = 0
        for file in self.__files:
            if file.file_status == 'Success':
                flag = 1
                break
        if (not self.__files) or (not flag):
            QMessageBox(QMessageBox.Warning, '警告', '请先转写语音文件').exec()
            return
        file_path, file_type = QFileDialog.getSaveFileName(self,
                                                           "文件保存",
                                                           os.getcwd(), "文档格式 (*.docx);;表格格式 (*.xlsx)")
        if not file_path:
            return
        if re.match('.*\.docx', file_type):
            self.write2doc(file_path)
        else:
            self.write2excel(file_path)

    # 把内容写为docx
    def write2doc(self, path):
        doc = Document()
        doc.styles['Normal'].font.name = u'宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        for file in self.__files:
            if file.file_status == 'Success':
                doc.add_paragraph(file.file_name)
                doc.add_paragraph(file.get_file_txt())
                doc.add_paragraph('\n')

        doc.save(path)

    # 把内容写为xlsx
    def write2excel(self, path):
        file_name =[]
        trans_content = []
        for file in self.__files:
            if file.file_status == 'Success':
                file_name.append(file.file_name)
                trans_content.append(file.get_file_txt())

        struct = {
            '文件名': file_name,
            '转写内容':trans_content
        }
        pd = DataFrame(struct)
        pd.to_excel(path)
