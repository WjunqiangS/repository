from PyQt5.QtWidgets import *
from voice_transform import VoiceTransThread
from docx import Document
from docx.oxml.ns import qn
from pandas import DataFrame
import os
import re


class TransformGui(QWidget):
    """ 本类包含了所有功能按钮以及功能按钮的槽函数
        __self_upload：把语音文件上传转写后返回语音对应的文本文件
    """
    def __init__(self):
        super(TransformGui, self).__init__()
        self.voice_trans_thread = None
        self.__files = []
        self.__init_control()

    def __init_control(self):
        # 创建功能按键
        self.__btn_upload = QPushButton('语音转写')
        self.__btn_save = QPushButton('保存修改内容')
        self.__btn_export = QPushButton('导出转写内容')
        self.__btn_save.setEnabled(False)

        # 绑定功能按钮的槽函数
        self.__btn_upload.clicked.connect(self.on_btn_upload_clicked)
        self.__btn_save.clicked.connect(self.on_btn_save_clicked)
        self.__btn_export.clicked.connect(self.on_btn_export_clicked)

        # 创建文本区
        self.__text = QPlainTextEdit()
        self.__text.textChanged.connect(self.on_text_changed)

        # 创建布局管理器，管理功能按键
        vlayout = QVBoxLayout()

        # 添加控件到布局管理器
        vlayout.addWidget(QLabel('语音转写区:'))
        vlayout.addWidget(self.__text)
        vlayout.addWidget(self.__btn_save)
        vlayout.addWidget(self.__btn_upload)
        vlayout.addWidget(self.__btn_export)

        self.setLayout(vlayout)

    # 语音上传转写功能按钮槽函数
    def on_btn_upload_clicked(self):
        if not self.__files:
            QMessageBox(QMessageBox.Warning, '警告', '请打开要转写的语音文件').exec()
        else:
            self.__text.clear()
            self.__text.appendPlainText('正在转写中...')
            self.__btn_upload.setEnabled(False)
            self.__text.setEnabled(False)
            self.__btn_save.setEnabled(False)
            self.voice_trans_thread = VoiceTransThread(self.__files)
            self.voice_trans_thread.trans_end.connect(self.voice_trans_end)
            for file in self.__files:
                file.transforming = True
            self.voice_trans_thread.start()

    # 获取打开的文件
    def get_files(self, file):
        # 如果发现语音转写线程被开启，则把后面加上文件状态改为正在转写状态
        if self.voice_trans_thread :
            file.transforming = True
        self.__files.append(file)

    # 语音转写完成之后退出线程
    def voice_trans_end(self, files):
        # 转写完成之后更新文件的信息
        self.__files = files

        # 设置各个控件的状态
        self.__text.clear()
        self.__text.appendPlainText('语音转写完成')
        self.__text.setEnabled(True)
        self.__btn_upload.setEnabled(True)
        self.__btn_save.setEnabled(False)

        # 退出线程
        self.voice_trans_thread.quit()
        self.voice_trans_thread = None

    # 点击文件列表的时候，显示选中文件的内容
    def show_file_txt(self, file_name):
        for file in self.__files:
            if file_name == file.file_name:
                break;
        # 获取被选中的文件
        self.cur_file = file
        self.__text.clear()
        # 如果文件完成转写，则显示转写的内容，如果还在转写，则显示正在转写中
        if file.finish_transform:
            self.__text.appendPlainText(file.get_file_txt())
        elif file.transforming:
            self.__text.appendPlainText('正在转写中...')
        self.__btn_save.setEnabled(False)

    # 保存修改的文本内容
    def on_btn_save_clicked(self):
        ret = QMessageBox.information(self, '警告', '是否保存文件', QMessageBox.No|QMessageBox.Yes)
        if ret == QMessageBox.Yes:
            self.cur_file.set_file_txt(self.__text.toPlainText())

    # 导出语音转写内容
    def on_btn_export_clicked(self):
        flag = 0
        for file in self.__files:
            if file.finish_transform:
                flag = 1
                break
        if (not self.__files) or (not flag):
            QMessageBox(QMessageBox.Warning, '警告', '请先转写语音文件').exec()
            return
        file_path, file_type = QFileDialog.getSaveFileName(self,
                                                           "文件保存",
                                                           os.getcwd(), "文档格式 (*.docx);;表格格式 (*.xlsx)")
        if not file_path:
            return
        if re.match('.*\.docx', file_type):
            self.write2doc(file_path)
        else:
            self.write2excel(file_path)

    # 转写内容被修改之后打开保存修改内容按钮
    def on_text_changed(self):
        self.__btn_save.setEnabled(True)

    # 把内容写为docx
    def write2doc(self, path):
        doc = Document()
        doc.styles['Normal'].font.name = u'宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        for file in self.__files:
            if file.finish_transform:
                doc.add_paragraph(file.file_name)
                doc.add_paragraph(file.get_file_txt())
                doc.add_paragraph('\n')

        doc.save(path)

    # 把内容写为xlsx
    def write2excel(self, path):
        file_name =[]
        trans_content = []
        for file in self.__files:
            if file.finish_transform:
                file_name.append(file.file_name)
                trans_content.append(file.get_file_txt())

        struct = {
            '文件名': file_name,
            '转写内容':trans_content
        }
        pd = DataFrame(struct)
        pd.to_excel(path)
